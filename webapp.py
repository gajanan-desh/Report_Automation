from flask import Flask, request, render_template, send_file, flash, redirect, url_for
import os
import tempfile
from docx.shared import Pt
import csv
from werkzeug.utils import secure_filename
import logging
from datetime import datetime
from script5 import main as generate_document, validate_csv_columns, required_fields
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document

def enforce_font_on_appended_content(doc, font_name="Helvetica", font_size=10.5, from_paragraph_idx=None):
    """
    Enforce font size and font family starting from a specific paragraph index,
    assuming everything before that is from the cover page.
    """
    logger.info(f"Enforcing font settings from paragraph index {from_paragraph_idx}")
    try:
        if from_paragraph_idx is None:
            from_paragraph_idx = 0  # default to all paragraphs

        for i, para in enumerate(doc.paragraphs):
            if i >= from_paragraph_idx:
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
    except Exception as e:
        logger.error(f"Failed to enforce font on appended content: {str(e)}")   
# Set up logging
log_directory = "logs"
if not os.path.exists(log_directory):
    os.makedirs(log_directory)

log_filename = os.path.join(log_directory, f"webapp_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log")
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_filename),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = 'your_secret_key_here' 
app.config['UPLOAD_FOLDER'] = tempfile.mkdtemp()
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024 

ALLOWED_EXTENSIONS = {'csv'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def merge_documents(first_doc_path, second_doc_path, output_path):
    """Merge two Word documents into one"""
    logger.info("Merging documents")
    try:
        # Open the first document (cover)
        first_doc = Document(first_doc_path)

        # Open the second document (VA report)
        second_doc = Document(second_doc_path)

        # Record how many paragraphs are in the cover page (to know where appended content starts)
        cover_paragraph_count = len(first_doc.paragraphs)

        # Append all elements from the second document to the first
        for element in second_doc.element.body:
            first_doc.element.body.append(element)

        # Apply font only to the newly appended content
        enforce_font_on_appended_content(first_doc, font_name="Helvetica", font_size=10.5, from_paragraph_idx=cover_paragraph_count)

        # Save the merged document
        first_doc.save(output_path)
        logger.info("Documents merged successfully")
        return True
    except Exception as e:
        logger.error(f"Error merging documents: {str(e)}")
        return False


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Get client information
        company_name = request.form.get('companyName', 'Company Name')
        network_type = request.form.get('networkType', 'External Network')
        assessment_date = request.form.get('assessmentDate')
        finding_id_prefix = request.form.get('findingIdPrefix')
        if assessment_date:
            # Convert from YYYY-MM-DD to your preferred format (e.g., DD/MM/YYYY)
            date_obj = datetime.strptime(assessment_date, '%Y-%m-%d')
            assessment_date = date_obj.strftime('%d/%m/%Y')
        else:
            assessment_date = datetime.now().strftime('%d/%m/%Y')
        findings_count = request.form.get('findingsCount', '5')

        if not finding_id_prefix:
            flash('Finding ID prefix is required')
            return redirect(request.url)
        
        logger.info(f"Processing report for {company_name}, {network_type}")
        
      
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)
        
        file = request.files['file']
        
       
        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)
        
        if file and allowed_file(file.filename):
            try:
                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                logger.info(f"File uploaded: {filename}")
                
            
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as csv_file:
                    csv_reader = csv.reader(csv_file)
                    header = next(csv_reader)
                    
                  
                    column_map = {}
                    for i, column_name in enumerate(header):
                        column_map[column_name] = i
                    
                    valid, error_message = validate_csv_columns(column_map, required_fields)
                    if not valid:
                        flash(f"CSV validation failed: {error_message}")
                        logger.error(f"CSV validation failed: {error_message}")
                        return redirect(request.url)
                
                # Save as dataset.csv for the script
                dataset_path = os.path.join(app.config['UPLOAD_FOLDER'], 'dataset.csv')
                if os.path.exists(dataset_path):
                    os.remove(dataset_path)  # Delete the existing file first
                os.rename(filepath, dataset_path)
                
                # Step 1: Generate the cover page document using DocxTemplate
                try:
                    template_path = os.path.join('templates', 'report_template.docx')
                    
                    # Make sure the template exists
                    if not os.path.exists(template_path):
                        flash("Report template not found")
                        logger.error("Report template not found")
                        return redirect(request.url)
                        
                    doc = DocxTemplate(template_path)
                    
                    context = {
                        'companyName': company_name,
                        'networkType': network_type,
                        'assessmentDate': assessment_date,
                        'findingsCount': findings_count,
                    }
                    
                    doc.render(context)
                    
                    # Save the cover document
                    cover_path = os.path.join(app.config['UPLOAD_FOLDER'], 'cover_document.docx')
                    doc.save(cover_path)
                    logger.info("Cover page generated successfully")
                    
                except Exception as e:
                    flash(f"Failed to generate cover page: {str(e)}")
                    logger.error(f"Failed to generate cover page: {str(e)}")
                    return redirect(request.url)
                
                # Step 2: Set the current working directory to the upload folder and run the VA report generation
                original_dir = os.getcwd()
                os.chdir(app.config['UPLOAD_FOLDER'])
                
                # Run the document generation
                success = generate_document(finding_id_prefix)
                
                # Reset working directory
                os.chdir(original_dir)
                
                if success:
                    output_path = os.path.join(app.config['UPLOAD_FOLDER'], 'output_document.docx')
                    if os.path.exists(output_path):
                        logger.info("VA report generated successfully")
                        
                        # Step 3: Merge the documents
                        merged_path = os.path.join(app.config['UPLOAD_FOLDER'], 'merged_report.docx')
                        merge_success = merge_documents(cover_path, output_path, merged_path)
                        
                        if merge_success:
                            logger.info("Final report generated successfully")
                            return send_file(
                                merged_path,
                                as_attachment=True,
                                download_name=f'{company_name}_vulnerability_report.docx',
                                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                            )
                        else:
                            flash("Failed to merge documents")
                            logger.error("Document merging failed")
                            # Fallback: return just the VA report
                            return send_file(
                                output_path,
                                as_attachment=True,
                                download_name='vulnerability_report.docx',
                                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                            )
                    else:
                        flash("Document was not generated properly")
                        logger.error("Document file not found after generation")
                else:
                    flash("Failed to generate the vulnerability report. Please check the logs for details.")
                    logger.error("VA report generation failed")
            
            except Exception as e:
                flash(f"An error occurred: {str(e)}")
                logger.exception("Exception during file processing")
                return redirect(request.url)
        else:
            flash('Only CSV files are allowed')
    
    # Get the required fields to display in the template
    required_columns = list(required_fields.values())
    
    return render_template('index.html', required_columns=required_columns)

if __name__ == '__main__':
    # Ensure templates directory exists
    if not os.path.exists('templates'):
        os.makedirs('templates')
    
    # Check if the template exists, if not create a placeholder
    template_path = os.path.join('templates', 'report_template.docx')
    if not os.path.exists(template_path):
        logger.warning("Report template not found. Please add a template file.")
        flash("Report template not found. Please add a template file.")
    
    logger.info("Starting web application")
    app.run