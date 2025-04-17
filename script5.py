from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import csv
import logging
import re
import os
import sys
from datetime import datetime
required_fields = {
            'name': 'Name',
            'description': 'Description',
            'cvs_score': 'CVSS v3.0 Base Score',
            'risk_factor': 'Risk Factor',
            'host': 'Host',
            'port': 'Port',
            'mitigation': 'Solution',
            'references': 'See Also'
        }
def validate_csv_columns(column_map, required_fields):
    """Validate if all required columns exist in the CSV file"""
    logger.info("Validating CSV columns")
    missing_fields = []
    for field, column_name in required_fields.items():
        if column_name not in column_map:
            missing_fields.append(column_name)
    
    if missing_fields:
        error_msg = f"Missing required columns in CSV: {', '.join(missing_fields)}"
        logger.error(error_msg)
        return False, error_msg
    
    logger.info("All required columns found in CSV")
    return True, ""
# Set up logging
log_directory = "logs"
if not os.path.exists(log_directory):
    os.makedirs(log_directory)

log_filename = os.path.join(log_directory, f"document_creation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log")
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_filename),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

def set_table_border_color(table, border_color="9C9C9C"):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._element.xpath('w:tblPr')[0]
    borders = OxmlElement('w:tblBorders')
    for border_position in ['top', 'start', 'bottom', 'end', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_position}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # width
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), border_color)
        borders.append(border)
    tbl_pr.append(borders)


def set_document_font(doc, font_name="Helvetica", font_size=11):
    """Set the default font for the entire document."""
    logger.info(f"Setting document font to {font_name}, size {font_size}")
    try:
        style = doc.styles['Normal']
        style.font.name = font_name
        style.font.size = Pt(font_size)
        
        # Ensure the correct font is applied to tables too
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
        logger.info("Document font set successfully")
    except Exception as e:
        logger.error(f"Failed to set document font: {str(e)}")
        raise

def make_cell_text_bold(cell):
    """Helper function to make cell text bold"""
    try:
        paragraph = cell.paragraphs[0]
        run = paragraph.runs
        if not run:  # If there's no run, create one
            run = paragraph.add_run(cell.text)
            run.bold = True
            cell.text = ''  # Clear the cell text since we're using the run
        else:
            run[0].bold = True
    except Exception as e:
        logger.error(f"Failed to make cell text bold: {str(e)}")
        raise

def format_affected_resources(table, affected_resources):
    logger.info(f"Formatting affected resources table for {len(affected_resources)} resources")
    try:
        cell = table.cell(1, 0)
        cell.text = ""  # Clear existing content
        
        # Sort the resources for consistent display
        resources = sorted(affected_resources)
        total = len(resources)
        
        # Clear all existing paragraphs in the cell except the first one
        for p in cell.paragraphs[1:]:
            p._element.getparent().remove(p._element)
        
        # Get the first paragraph
        first_paragraph = cell.paragraphs[0]
        first_paragraph.text = ""  # Ensure it's empty
        
        # For fewer than 5 resources: vertical layout (single column)
        if total < 5:
            # Add first resource to the existing paragraph
            if resources:
                run = first_paragraph.add_run(resources[0])
                run.font.size = Pt(10.5)
            
            # Add remaining resources in separate paragraphs
            for i in range(1, total):
                paragraph = cell.add_paragraph()
                run = paragraph.add_run(resources[i])
                run.font.size = Pt(10.5)
                
                # Set paragraph spacing
                paragraph.space_after = Pt(0)
                paragraph.space_before = Pt(0)
        
        # For 5 or more resources: two-column layout
        else:
            # Calculate number of rows needed (ceiling division)
            rows = (total + 1) // 2  # This ensures we round up
            
            # Add first row to existing paragraph
            if resources:
                run = first_paragraph.add_run(resources[0])
                run.font.size = Pt(10)
                padding = 35 - len(resources[0])  # Adjust 35 based on your needs
                run = first_paragraph.add_run(" " * padding)
                
                # Add right column for first row if it exists
                if rows < len(resources):
                    run = first_paragraph.add_run(resources[rows])
                    run.font.size = Pt(10)
            
            # Set paragraph spacing
            first_paragraph.space_after = Pt(0)
            first_paragraph.space_before = Pt(0)
            
            # Create paragraph for each remaining row (starting from index 1)
            for i in range(1, rows):
                paragraph = cell.add_paragraph()
                
                # Add left column
                if i < len(resources):
                    run = paragraph.add_run(resources[i])
                    run.font.size = Pt(10)  
                    padding = 35 - len(resources[i])  # Adjust based on your needs
                    run = paragraph.add_run(" " * padding)
                    
                    # Add right column if it exists
                    if i + rows < len(resources):
                        run = paragraph.add_run(resources[i + rows])
                        run.font.size = Pt(10)
                
                # Set paragraph spacing
                paragraph.space_after = Pt(0)
                paragraph.space_before = Pt(0)
        logger.info("Successfully formatted affected resources")
    except Exception as e:
        logger.error(f"Failed to format affected resources: {str(e)}")
        raise
def create_summary_table(doc, grouped_vulnerabilities):
    """Create a summary table of all vulnerabilities at the beginning of the document"""
    logger.info("Creating vulnerability summary table")
    try:
        
        
        # Create a table with appropriate columns
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Set the header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "#"
        header_cells[1].text = "Title"
        header_cells[2].text = "Risk"
        header_cells[3].text = "CVSS"
        header_cells[4].text = "Finding ID"
        
        # Make header row bold, centered, and apply shading - using #5B9BD5 for header
        for i, cell in enumerate(header_cells):
            # Center align all header cells
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Make text bold
            cell.paragraphs[0].runs[0].bold = True
            
            # Set font to Helvetica and size to 10
            cell.paragraphs[0].runs[0].font.name = "Helvetica"
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            # Apply shading
            tcPr = cell._tc.get_or_add_tcPr()
            tcShading = OxmlElement('w:shd')
            tcShading.set(qn('w:fill'), '5B9BD5')  # Blue header (#5B9BD5)
            tcShading.set(qn('w:val'), 'clear')
            tcPr.append(tcShading)
            
            # Text color white for header
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Change the table borders to #5B9BD5
        tblPr = table._element.xpath('w:tblPr')[0]
        tblBorders = OxmlElement('w:tblBorders')
        
        # Define all borders with the same color
        for border_position in ['top', 'start', 'bottom', 'end', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_position}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Border width
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '5B9BD5')  # Border color (#5B9BD5)
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
        
        # Sort vulnerabilities by risk factor priority and CVSS score
        risk_priority = {"critical": 1, "high": 2, "medium": 3, "low": 4, "informational": 5}
        
        # Helper function to safely convert CVSS to float
        def safe_cvss_float(cvss):
            try:
                return -float(cvss) if cvss and cvss.strip() else 0
            except (ValueError, TypeError):
                return 0
        
        sorted_vulns = sorted(
            grouped_vulnerabilities.values(), 
            key=lambda x: (risk_priority.get(x["risk_factor"].lower(), 999), safe_cvss_float(x["cvs_score"]))
        )
        
        # Add vulnerability data with alternating row colors
        for i, vuln in enumerate(sorted_vulns):
            row = table.add_row()
            
            # Apply alternating row colors
            if i % 2 == 1:  # Odd rows (0-indexed, so row 1, 3, 5, etc.)
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcShading = OxmlElement('w:shd')
                    tcShading.set(qn('w:fill'), 'DEEAF6')  # Light blue background (#DEEAF6)
                    tcShading.set(qn('w:val'), 'clear')
                    tcPr.append(tcShading)
            
            # Fill row data
            row.cells[0].text = str(i + 1)  # Sequential numbering
            row.cells[1].text = vuln["name"]
            row.cells[2].text = vuln["risk_factor"].capitalize()
            row.cells[3].text = str(vuln["cvs_score"]) if vuln["cvs_score"] else "-"
            row.cells[4].text = vuln["finding_id"]
            
            # Apply formatting to all cells
            for cell_idx, cell in enumerate(row.cells):
                # Set font to Helvetica and size to 10 for all ceBUlls
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Helvetica"
                        run.font.size = Pt(10)

                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # Center align specific columns: #, Risk, and CVSS
                if cell_idx in [0, 2, 3]:  # #, Risk, and CVSS columns
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Color coding for risk cells
            risk_colors = {
                "critical": "C00000",  # Red
                "high": "FFC000",      # Orange
                "medium": "FFFF00",    # Yellow
                "low": "92D050",       # Green
                "informational": "9BC2E6"  # Light blue
            }
            
            risk_cell = row.cells[2]
            risk_factor = vuln["risk_factor"].lower()
            
            if risk_factor in risk_colors:
                # Apply cell shading
                tcPr = risk_cell._tc.get_or_add_tcPr()
                tcShading = OxmlElement('w:shd')
                tcShading.set(qn('w:fill'), risk_colors[risk_factor])
                tcShading.set(qn('w:val'), 'clear')
                tcPr.append(tcShading)
        
        # Set column widths
        table.autofit = False
        table.allow_autofit = False
        
        # Set column widths (adjust as needed)
        widths = [Pt(30), Pt(250), Pt(80), Pt(60), Pt(80)]
        for i, width in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = width
        
        # Add space after the table
        doc.add_paragraph()
        
        logger.info("Summary table created successfully")
        return True
    except Exception as e:
        logger.error(f"Failed to create summary table: {str(e)}")
        # Continue execution even if summary table fails
        return False

def format_cvss_score(score):
    """Format CVSS score to always show at least one decimal place"""
    try:
        # Try to convert to float
        float_score = float(score)
        # Format with one decimal place
        return f"{float_score:.1f}"
    except (ValueError, TypeError):
        # If conversion fails, return the original score
        return score
       
def create_table(doc, heading):
    logger.info(f"Creating table with heading: {heading}")
    try:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(heading)
        run.bold = True
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(0)
        # Apply the theme color (matches Word's default heading color)
        rPr = run._element.get_or_add_rPr()
        color = parse_xml(r'<w:color {} w:val="365F91"/>'.format(nsdecls('w')))  # Blue theme color
        rPr.append(color)

        # Row 0-1
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        set_table_border_color(table)
        table.cell(0, 0).width = Pt(150)
        table.cell(1, 0).width = Pt(150)
        table.cell(0, 1).width = Pt(350)
        table.cell(1, 1).width = Pt(350)
        
        # Set headers and make bold
        cells = [
            (0, 0, 'Finding ID'),
            (0, 1, 'Description')
        ]
        
        for row, col, text in cells:
            cell = table.cell(row, col)
            cell.text = text
            make_cell_text_bold(cell)
        
        table.cell(1, 0).text = ''
        table.cell(1, 1).text = ''

        # Apply shading
        for col in range(2):
            table_header = table.cell(0, col)
            tcPr = table_header._tc.get_or_add_tcPr()
            tcShading = OxmlElement('w:shd')
            tcShading.set(qn('w:fill'), 'C0D4EC')
            tcShading.set(qn('w:themeTint'), '40')
            tcPr.append(tcShading)

        # Row 2-3
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        set_table_border_color(table)
        # Set cell widths
        for col in range(3):
            table.cell(0, col).width = Pt(150)
            table.cell(1, col).width = Pt(150)

        # Set headers and make bold
        cells = [
            (0, 0, 'CVS Score'),
            (0, 1, 'Risk Rating'),
            (0, 2, 'Remote Exploitability')
        ]
        
        for row, col, text in cells:
            cell = table.cell(row, col)
            cell.text = text
            make_cell_text_bold(cell)

        # Clear data cells
        for col in range(3):
            table.cell(1, col).text = ''

        # Apply shading
        for col in range(3):
            table_header = table.cell(0, col)
            tcPr = table_header._tc.get_or_add_tcPr()
            tcShading = OxmlElement('w:shd')
            tcShading.set(qn('w:fill'), 'C0D4EC')
            tcShading.set(qn('w:themeTint'), '40')
            tcPr.append(tcShading)

        # Row 4-5
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        set_table_border_color(table)

        
        # Set cell widths
        for row in table.rows:
            for cell in row.cells:
                cell.width = Pt(150)
        # Merge column 0 and column 1 to create space for "Affected Resource"
        table.cell(0, 0).merge(table.cell(0, 1))  # Merge first row, columns 0 and 1
        table.cell(1, 0).merge(table.cell(1, 1))  # Merge second row, columns 0 and 1
        # Set headers and make bold
        cells = [
            (0, 0, 'Affected Resource'),
            (0, 2, 'Module Name')
        ]
        
        for row, col, text in cells:
            cell = table.cell(row, col)
            cell.text = text
            make_cell_text_bold(cell)

        # Clear data cells
        table.cell(1, 0).text = ''
        table.cell(1, 1).text = ''

        # Apply shading
        for col in range(3):
            table_header = table.cell(0, col)
            tcPr = table_header._tc.get_or_add_tcPr()
            tcShading = OxmlElement('w:shd')
            tcShading.set(qn('w:fill'), 'C0D4EC')
            tcShading.set(qn('w:themeTint'), '40')
            tcPr.append(tcShading)

        # Row 6-7
        table = doc.add_table(rows=2, cols=1)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        set_table_border_color(table)

        # Set cell widths
        for row in table.rows:
            for cell in row.cells:
                cell.width = Pt(300)

        # Set header and make bold
        cell = table.cell(0, 0)
        cell.text = 'Security Risk'
        make_cell_text_bold(cell)
        
        # Clear data cell
        table.cell(1, 0).text = ''

        # Apply shading
        table_header = table.cell(0, 0)
        tcPr = table_header._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), 'C0D4EC')
        tcShading.set(qn('w:themeTint'), '40')
        tcPr.append(tcShading)

        # Row 8
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        set_table_border_color(table)

        # Set cell widths
        for row in table.rows:
            for cell in row.cells:
                cell.width = Pt(150)

        # Set header and make bold
        cell = table.cell(0, 0)
        cell.text = 'Business Impact'
        make_cell_text_bold(cell)
        cell_paragraph = cell.paragraphs[0]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Also center the text vertically in the cell
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Clear data cell
        table.cell(0, 1).text = ''

        # Row 9-10
        table = doc.add_table(rows=2, cols=1)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        set_table_border_color(table)
        # Set cell widths
        for row in table.rows:
            for cell in row.cells:
                cell.width = Pt(300)

        # Set header and make bold
        cell = table.cell(0, 0)
        cell.text = 'Workaround / Mitigation'
        make_cell_text_bold(cell)
        
        # Clear data cell
        table.cell(1, 0).text = ''

        # Apply shading
        table_header = table.cell(0, 0)
        tcPr = table_header._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), 'C0D4EC')
        tcShading.set(qn('w:themeTint'), '40')
        tcPr.append(tcShading)

        # Row 11-12
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        set_table_border_color(table)
        # Set cell widths
        for row in table.rows:
            for cell in row.cells:
                cell.width = Pt(150)
        table.cell(0,0).merge(table.cell(0,1))
        table.cell(1, 0).merge(table.cell(1,1)) 
        # Set headers and make bold
        cells = [
            (0, 0, 'Tool used'),
            (0, 2, 'References')
        ]
        
        for row, col, text in cells:
            cell = table.cell(row, col)
            cell.text = text
            make_cell_text_bold(cell)

        # Clear data cells
        table.cell(1, 0).text = ''
        table.cell(1, 1).text = ''

        # Apply shading
        for col in range(3):
            table_header = table.cell(0, col)
            tcPr = table_header._tc.get_or_add_tcPr()
            tcShading = OxmlElement('w:shd')
            tcShading.set(qn('w:fill'), 'C0D4EC')
            tcShading.set(qn('w:themeTint'), '40')
            tcPr.append(tcShading)

        # Row 13-14
        table = doc.add_table(rows=2, cols=1)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        set_table_border_color(table)

        # Set cell widths
        for row in table.rows:
            for cell in row.cells:
                cell.width = Pt(300)

        # Set header and make bold
        cell = table.cell(0, 0)
        cell.text = 'Proof of Concept (POC)'
        make_cell_text_bold(cell)
        
        # Clear data cell
        table.cell(1, 0).text = ''

        # Apply shading
        table_header = table.cell(0, 0)
        tcPr = table_header._tc.get_or_add_tcPr()
        tcShading = OxmlElement('w:shd')
        tcShading.set(qn('w:fill'), 'C0D4EC')
        tcShading.set(qn('w:themeTint'), '40')
        tcPr.append(tcShading)
        
        logger.info(f"Table created successfully for heading: {heading}")
        return True
    except Exception as e:
        logger.error(f"Failed to create table for {heading}: {str(e)}")
        raise

KEYWORDS = {
    'Apache': 'Apache',
    'Window': 'Windows',
    'SSH': 'SSH',
    'Oracle': 'Oracle',
    'DNS': 'DNS',
}
def determine_risk_factor(cvss_score):
    """Determine risk factor based on CVSS score ranges"""
    try:
        score = float(cvss_score)
        
        if score >= 9.0 and score <= 10.0:
            return "critical"
        elif score >= 7.0 and score < 9.0:
            return "high"
        elif score >= 4.0 and score < 7.0:
            return "medium"
        elif score > 0.0 and score < 4.0:
            return "low"
        elif score == 0.0:
            return "informational"
        else:
            logger.warning(f"CVSS score {score} is outside expected range (0-10)")
            return "informational"  # Default to lowest risk level
    except (ValueError, TypeError):
        logger.error(f"Invalid CVSS score: {cvss_score}")
        return "informational"  # Default to lowest risk level

def get_module_name(name, predefined_keywords):
    """Helper function to match module names from predefined keywords"""
    logger.debug(f"Finding module name for: {name}")
    try:
        # Iterate through the predefined keywords and check if they match the name
        for keyword in predefined_keywords:
            if keyword.lower() in name.lower():
                logger.debug(f"Found matching module name: {keyword}")
                return keyword
        logger.debug("No matching module name found")
        return ""  # Return empty string if no match is found
    except Exception as e:
        logger.error(f"Error in get_module_name: {str(e)}")
        return ""

def extract_description(text):
    """
    Extract a meaningful description from text with better handling of:
    1. Periods in version numbers (like 1.2.3.4)
    2. Periods in product names (like ASP.NET)
    3. Periods followed by alphanumeric characters
    
    The function aims to preserve technical information while providing a concise description.
    """
    logger.debug(f"Extracting description from: {text}")
    
    # Handle empty or None text
    if not text:
        logger.debug("Empty text, returning as is")
        return text
    
    # Check if the text is fairly short (less than 200 chars) - return it as is
    if len(text) < 200:
        logger.debug("Text is short, returning full text")
        return text
    
    # Use regex to identify special patterns that should be preserved
    
    
    # Find all occurrences of patterns like "X.Y" where X and Y can be alphanumeric
    # This will match version numbers, product names with periods, etc.
    special_patterns = re.findall(r'\b\w+\.\w+\b', text)
    logger.debug(f"Found special patterns to preserve: {special_patterns}")
    
    # Split by sentences (periods followed by space and capital letter)
    # This is a more reliable way to identify actual sentence boundaries
    sentences = re.split(r'\.(?=\s+[A-Z])', text)
    
    # If there's only one sentence or no clear sentence breaks, return the whole text
    if len(sentences) <= 1:
        logger.debug("Single sentence or no clear sentence breaks, returning full text")
        return text
    
    # Take the first two sentences as the base description
    result = sentences[0] + '.'
    
    # Check if we need to add more context
    if len(sentences) > 1:
        # Add the second sentence if it's not too long
        if len(sentences[1]) < 150:
            result += sentences[1] + '.'
    
    # Check if any special patterns (like version numbers) were cut off
    for pattern in special_patterns:
        if pattern not in result and len(result) + len(pattern) + 20 < 400:  # Avoid making it too long
            # Add a note about the version/pattern if it was important
            result += f" Relevant: {pattern}."
    
    logger.debug(f"Extracted description: {result}")
    return result
def reset_document_fonts(doc, font_name="Helvetica", font_size=10.5):
    """Reset all fonts in the document to ensure consistency"""
    logger.info(f"Resetting document fonts to {font_name}")
    
    # Set default style
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)
    
    # Go through all paragraphs in the document
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
    
    # Go through all tables in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
    
    logger.info("Document fonts reset successfully")

def append_data(doc, csv_row_no, data_to_append, predefined_keywords):
    logger.info(f"Appending data for vulnerability {data_to_append['finding_id']}")
    try:
        # Get total number of tables in document
        total_tables = len(doc.tables)
        logger.debug(f"Total tables in document: {total_tables}")
        
        # Calculate starting table index for this vulnerability
        # With summary table at the beginning, we need to adjust the index
        # Each vulnerability has 8 tables, and we need to account for the summary table
        table_start_index = 1 + (csv_row_no * 8)  # 1 for summary table
        
        # Check if we have enough tables
        if table_start_index + 7 >= total_tables:
            logger.error(f"Not enough tables in document. Expected index {table_start_index}, total tables: {total_tables}")
            raise IndexError(f"Table index out of range: {table_start_index} >= {total_tables}")
        
        def set_cell_text(cell, text_content):
            cell.text = text_content
            # Apply font to all runs in all paragraphs of the cell
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Helvetica"
                    run.font.size = Pt(10.5)

        # table - 0 (Finding ID and Description)
        finding_id_cell = doc.tables[table_start_index].rows[1].cells[0]
        finding_id_cell.text = data_to_append['finding_id']
        finding_id_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Prepend the required text to the description
        description = f"During Vulnerability assessment and Penetration testing we observed that, {data_to_append['description']}"
        doc.tables[table_start_index].rows[1].cells[1].text = description

        # table - 1 (CVSS Score and Risk Factor)
        doc.tables[table_start_index + 1].rows[1].cells[0].text = format_cvss_score(data_to_append['cvs_score'])
        
        cell = doc.tables[table_start_index + 1].rows[1].cells[1]
        cell.text = data_to_append['risk_factor'].capitalize()

        risk_colors = {
            "critical": "C00000",
            "high": "FFC000",
            "medium": "FFFF00",
            "low": "92D050",
            "informational": "9BC2E6",
        }
        
        risk_color = risk_colors.get(data_to_append['risk_factor'].lower())
        if risk_color:
            tcPr = cell._tc.get_or_add_tcPr()
            tcShading = OxmlElement('w:shd')
            tcShading.set(qn('w:fill'), risk_color)
            tcPr.append(tcShading)

        # remote_exploitability
        doc.tables[table_start_index + 1].rows[1].cells[2].text = "Yes"

        # table - 2 (Affected Resource & Module Name)
        table = doc.tables[table_start_index + 2]
        
        # Get affected resources as a list
        affected_resources = data_to_append["affected_resource"].split("\n")
        
        # Use new formatting function
        format_affected_resources(table, affected_resources)

        # Module Name - dynamically matched from predefined keywords
        module_name_cell = table.cell(1, 2)
        module_name_cell.text = get_module_name(data_to_append['name'], predefined_keywords)
        module_name_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # table - 3 (Security Risk)
        doc.tables[table_start_index + 3].rows[1].cells[0].text = ""

        # table - 4 (Business Impact)
        doc.tables[table_start_index + 4].rows[0].cells[1].text = ""

        # table - 5 (Workaround / Mitigation)
        solution = f"It is recommended: \n-To {data_to_append['mitigation']}"
        doc.tables[table_start_index + 5].rows[1].cells[0].text = solution

        # table - 6 (Tool Used & References)
        doc.tables[table_start_index + 6].rows[1].cells[0].text = "Nessus"
        tool_cell = doc.tables[table_start_index + 6].rows[1].cells[0]
        tool_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        doc.tables[table_start_index + 6].rows[1].cells[2].text = data_to_append['references']

        # table - 7 (Proof of Concept)
        proof_of_concept_cell = doc.tables[table_start_index + 7].rows[1].cells[0]
        existing_text = proof_of_concept_cell.text  # Preserve existing content
        proof_of_concept_cell.text = ""  # Clear the cell to format text properly

        # Adding "Figure 1 - Shows" with "Figure 1" in bold
        proof_paragraph = proof_of_concept_cell.paragraphs[0]
        run = proof_paragraph.add_run("\nFigure 1")
        run.bold = True
        proof_paragraph.add_run(" - Shows " + existing_text)
        
        logger.info(f"Successfully appended data for vulnerability {data_to_append['finding_id']}")
        return True
    except Exception as e:
        logger.error(f"Failed to append data for {data_to_append['finding_id']}: {str(e)}")
        raise


def main(finding_id_prefix=None):
    logger.info("Starting document creation process")
    
    try:
        # Create a new document
        doc = Document()
        
        logger.info("Created new document")
        
        set_document_font(doc, "Helvetica", 10.5)
        grouped_vulnerabilities = {}  

        # Add a title
        
        
        csv_file_path = 'dataset.csv'
        logger.info(f"Using CSV file: {csv_file_path}")

       
        if not os.path.exists(csv_file_path):
            logger.error(f"CSV file not found: {csv_file_path}")
            raise FileNotFoundError(f"CSV file not found: {csv_file_path}")

        
        finding_id_counter = 1

        try:
            with open(csv_file_path, 'r', encoding='utf-8', errors='ignore') as csv_file:
                logger.info("Successfully opened CSV file")
                
                
                try:
                    csv_reader = csv.reader(csv_file)
                    
                    # Read the header row to determine column indices
                    header = next(csv_reader)
                    logger.info(f"CSV header read successfully with {len(header)} columns")
                    
                    # Create a mapping of column names to indices
                    column_map = {}
                    for i, column_name in enumerate(header):
                        column_map[column_name] = i
                    
                    # Validate CSV columns
                    valid, error_message = validate_csv_columns(column_map, required_fields)
                    if not valid:
                        logger.error(f"CSV validation failed: {error_message}")
                        logger.info(f"Available columns: {', '.join(column_map.keys())}")
                        raise ValueError(error_message)
                    
                    # Process each row in the CSV
                    row_count = 0
                    for csv_row_data in csv_reader:
                        row_count += 1
                        logger.info(f"Processing row {row_count}")
                        
                        try:
                            # Extract data using column names instead of hardcoded indices
                            vulnerability_name = csv_row_data[column_map[required_fields['name']]]
                            host = csv_row_data[column_map[required_fields['host']]]
                            port = csv_row_data[column_map[required_fields['port']]]

                            # Check if port is 0 and format accordingly
                            if port == "0":
                                affected_host = host  # Use only the host when port is 0
                            else:
                                affected_host = f"{host}:{port}"  # Use host:port format for non-zero ports
                                                    
                            # Get other required fields
                            full_description = csv_row_data[column_map[required_fields['description']]]
                            description = extract_description(full_description)

                            cvs_score = csv_row_data[column_map[required_fields['cvs_score']]]

                            # Calculate risk factor based on CVSS score instead of getting from CSV
                            risk_factor = determine_risk_factor(cvs_score)
                            mitigation = csv_row_data[column_map[required_fields['mitigation']]]
                            
                            # Handle possible empty references
                            references_idx = column_map[required_fields['references']]
                            references = csv_row_data[references_idx].split('\n')[0] if csv_row_data[references_idx] else ""
                        
                            if vulnerability_name in grouped_vulnerabilities:
                                grouped_vulnerabilities[vulnerability_name]["affected_resource"].add(affected_host)
                            else:
                                # Assign a new finding ID and store all required details
                                grouped_vulnerabilities[vulnerability_name] = {
                                    "name": vulnerability_name,
                                    "finding_id": f"{finding_id_prefix}-{finding_id_counter:02d}",
                                    "description": description,
                                    "cvs_score": format_cvss_score(cvs_score),
                                    "risk_factor": risk_factor,
                                    "remote_exploitability": "Yes",
                                    "affected_resource": {affected_host},
                                    "mitigation": mitigation,
                                    "references": references,
                                }
                                finding_id_counter += 1  # Increment counter
                            
                            logger.info(f"Successfully processed vulnerability: {vulnerability_name}")
                        
                        except IndexError as e:
                            logger.error(f"Invalid data in row {row_count}: {str(e)}")
                            logger.warning(f"Skipping row {row_count} due to missing or invalid data")
                            continue
                        except Exception as e:
                            logger.error(f"Error processing row {row_count}: {str(e)}")
                            logger.warning(f"Skipping row {row_count} due to processing error")
                            continue
                
                except csv.Error as e:
                    logger.error(f"CSV parsing error: {str(e)}")
                    raise
        
        except FileNotFoundError:
            logger.error(f"CSV file not found: {csv_file_path}")
            raise
        except PermissionError:
            logger.error(f"Permission denied when accessing CSV file: {csv_file_path}")
            raise
        except Exception as e:
            logger.error(f"Unexpected error when reading CSV file: {str(e)}")
            raise

        # Check if we have any vulnerabilities to process
        if not grouped_vulnerabilities:
            logger.warning("No valid vulnerabilities found in the CSV file")
            raise ValueError("No valid vulnerabilities found in the CSV file")

        logger.info(f"Found {len(grouped_vulnerabilities)} unique vulnerabilities")
        create_summary_table(doc, grouped_vulnerabilities)

# Add a page break after the summary table
        doc.add_section(WD_SECTION.NEW_PAGE)
        # Create tables for each vulnerability
        for index, (vulnerability_name, data_to_append) in enumerate(grouped_vulnerabilities.items()):
            logger.info(f"Creating table for vulnerability {index + 1}: {vulnerability_name}")
            
            # Fix: Strip spaces before joining affected resources
            data_to_append["affected_resource"] = "\n".join(res.strip() for res in sorted(data_to_append["affected_resource"]))

            # Create the table
            numbered_finding_name = f"{index + 1}. {vulnerability_name}"
            create_table(doc, numbered_finding_name)
        
            # Add new page except for the last vulnerability
            if index < len(grouped_vulnerabilities) - 1:
                doc.add_section(WD_SECTION.NEW_PAGE)
                logger.info("Added new page for next vulnerability")
            
            append_data(doc, index, data_to_append, KEYWORDS)

        # Save the document
        try:
            
            doc.save('output_document.docx')
            logger.info("Document saved successfully as 'output_document.docx'")
        except PermissionError:
            logger.error("Permission denied when saving the document. Check if the file is open in another application.")
            raise
        except Exception as e:
            logger.error(f"Failed to save document: {str(e)}")
            raise

        logger.info("Document creation process completed successfully")
        return True

    except Exception as e:
        logger.critical(f"Document creation failed: {str(e)}")
        return False

if __name__ == "__main__":
    try:
        test_prefix = input("Enter finding ID prefix (e.g., ABC): ")
        if not test_prefix:
            print("Finding ID prefix is required")
            exit(1)
        
        success = main(test_prefix)
        success = main()
        if success:
            print("Document created successfully!")
            logger.info("Script execution completed successfully")
        else:
            print("Document creation failed. Check the log file for details.")
            logger.error("Script execution failed")
    except Exception as e:
        print(f"An unexpected error occurred: {str(e)}")
        logger.critical(f"Unhandled exception: {str(e)}", exc_info=True)